"""
utils/document_reader.py
PointList v0.14.25experiment
Motor de extracción de texto multiformato para el Chatbot PointBit.
Soporta: PDF, DOCX, TXT, MD, JSON, CSV, LOG, CODE (PY, JS, HTML, CSS, SQL, etc).
"""

import os
import zipfile
import xml.etree.ElementTree as ET

SUPPORTED_EXTENSIONS = [
    ".pdf", ".docx", ".txt", ".md", ".json", ".csv",
    ".log", ".py", ".js", ".html", ".css", ".xml",
    ".sql", ".yaml", ".yml", ".env"
]

def is_supported_document(file_path: str) -> bool:
    ext = os.path.splitext(file_path)[1].lower()
    return ext in SUPPORTED_EXTENSIONS

def extract_text_from_file(file_path: str, max_chars: int = 25000) -> dict:
    """
    Extrae el contenido de texto de un archivo según su extensión.
    Retorna un diccionario con estado, texto extraído, nombre de archivo, número de palabras y metadatos.
    """
    if not os.path.exists(file_path):
        return {"ok": False, "error": "El archivo especificado no existe en el disco."}

    ext = os.path.splitext(file_path)[1].lower()
    filename = os.path.basename(file_path)

    text_extracted = ""

    try:
        # 1. Documentos PDF
        if ext == ".pdf":
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                pages_text = []
                for idx, page in enumerate(reader.pages):
                    t = page.extract_text() or ""
                    if t.strip():
                        pages_text.append(f"--- Página {idx + 1} ---\n{t.strip()}")
                text_extracted = "\n\n".join(pages_text)
            except Exception as e:
                return {"ok": False, "error": f"No se pudo extraer el contenido del PDF '{filename}': {str(e)}"}

        # 2. Documentos Word (.docx)
        elif ext == ".docx":
            try:
                import docx
                doc = docx.Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        row_txt = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_txt:
                            full_text.append(row_txt)
                text_extracted = "\n".join(full_text)
            except Exception:
                # Fallback nativo usando zipfile y ElementTree (sin dependencias adicionales)
                try:
                    with zipfile.ZipFile(file_path) as z:
                        xml_content = z.read("word/document.xml")
                        tree = ET.fromstring(xml_content)
                        texts = [node.text for node in tree.iter() if node.text]
                        text_extracted = "\n".join(texts)
                except Exception as e:
                    return {"ok": False, "error": f"No se pudo leer el archivo DOCX '{filename}': {str(e)}"}

        # 3. Texto plano, código fuente, CSV, JSON, Markdown, etc.
        elif ext in [".txt", ".md", ".json", ".csv", ".log", ".py", ".js", ".html", ".css", ".xml", ".sql", ".yaml", ".yml", ".env"]:
            for encoding in ["utf-8", "latin-1", "cp1252", "ascii"]:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        text_extracted = f.read()
                    break
                except UnicodeDecodeError:
                    continue

        else:
            # Intento de lectura UTF-8 estándar
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    text_extracted = f.read()
            except Exception:
                return {"ok": False, "error": f"El formato de archivo '{ext}' no es compatible para extracción de texto."}

        if not text_extracted or not text_extracted.strip():
            return {"ok": False, "error": f"El documento '{filename}' está vacío o no contiene texto procesable."}

        truncated = False
        if len(text_extracted) > max_chars:
            text_extracted = text_extracted[:max_chars] + f"\n\n... [Texto truncado a {max_chars} caracteres para optimizar la velocidad del análisis]"
            truncated = True

        words_count = len(text_extracted.split())

        return {
            "ok": True,
            "filename": filename,
            "extension": ext,
            "text": text_extracted.strip(),
            "num_words": words_count,
            "truncated": truncated
        }

    except Exception as ex:
        return {"ok": False, "error": f"Error al procesar el archivo '{filename}': {str(ex)}"}
